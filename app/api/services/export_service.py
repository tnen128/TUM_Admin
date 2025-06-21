from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
import os
from datetime import datetime
from typing import Dict

class DocumentExporter:
    def __init__(self):
        self.tum_blue = (0, 101, 189)  # TUM Corporate Blue
        self.temp_dir = tempfile.gettempdir()

    def _create_filename(self, doc_type: str, extension: str) -> str:
        """Create a standardized filename with timestamp"""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_doc_type = doc_type.lower().replace(" ", "_")
        return f"TUM_{safe_doc_type}_{timestamp}.{extension}"

    def export_to_pdf(self, content: str, metadata: Dict[str, str]) -> str:
        """Export document to PDF with TUM formatting"""
        pdf = FPDF()
        pdf.add_page()
        
        # Add header
        pdf.set_font("Arial", "B", 16)
        pdf.set_text_color(*self.tum_blue)
        pdf.cell(0, 10, f"TUM {metadata.get('doc_type', 'Document')}", ln=True, align="C")
        
        # Add metadata
        pdf.set_font("Arial", "I", 10)
        pdf.set_text_color(128, 128, 128)
        pdf.cell(0, 10, f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True)
        pdf.cell(0, 10, f"Tone: {metadata.get('tone', 'Standard')}", ln=True)
        
        # Add content
        pdf.set_font("Arial", "", 12)
        pdf.set_text_color(0, 0, 0)
        pdf.multi_cell(0, 10, content)
        
        # Save to temp file
        filename = self._create_filename(metadata.get('doc_type', 'document'), "pdf")
        filepath = os.path.join(self.temp_dir, filename)
        pdf.output(filepath)
        
        return filepath

    def export_to_docx(self, content: str, metadata: Dict[str, str]) -> str:
        """Export document to DOCX with TUM formatting"""
        doc = Document()
        
        # Add header
        header = doc.add_heading(f"TUM {metadata.get('doc_type', 'Document')}", level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        doc.add_paragraph(f"Tone: {metadata.get('tone', 'Standard')}")
        
        # Add separator
        doc.add_paragraph("=" * 50)
        
        # Add content
        doc.add_paragraph(content)
        
        # Save to temp file
        filename = self._create_filename(metadata.get('doc_type', 'document'), "docx")
        filepath = os.path.join(self.temp_dir, filename)
        doc.save(filepath)
        
        return filepath

    def export_to_txt(self, content: str, metadata: Dict[str, str]) -> str:
        """Export document to plain text with minimal formatting"""
        filename = self._create_filename(metadata.get('doc_type', 'document'), "txt")
        filepath = os.path.join(self.temp_dir, filename)
        
        with open(filepath, 'w') as f:
            f.write(f"TUM {metadata.get('doc_type', 'Document')}\n")
            f.write("=" * 50 + "\n\n")
            f.write(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n")
            f.write(f"Tone: {metadata.get('tone', 'Standard')}\n")
            f.write("=" * 50 + "\n\n")
            f.write(content)
        
        return filepath

    def export_document(self, content: str, metadata: Dict[str, str], format: str) -> str:
        """Export document in the specified format"""
        if format == "pdf":
            return self.export_to_pdf(content, metadata)
        elif format == "docx":
            return self.export_to_docx(content, metadata)
        elif format == "txt":
            return self.export_to_txt(content, metadata)
        else:
            raise ValueError(f"Unsupported format: {format}") 